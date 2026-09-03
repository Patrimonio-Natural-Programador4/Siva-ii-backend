import re
import shutil
import subprocess
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

from docx import Document

from config import STORAGE_PATH
from dto.TermsReferenceDTO import TdrFormField

_PATRON_MARCADOR = re.compile(r"<<\s*([a-zA-Z0-9_]+)\s*>>")
_PATRON_HTML = re.compile(r"<[a-zA-Z][^>]*>")


def generar_tdr_documento(plantilla_nombre: str, campos: list[TdrFormField], identificador: str) -> Path:
    ruta_plantilla = Path(STORAGE_PATH, "soportes", "tdr", "template", plantilla_nombre)
    if not ruta_plantilla.exists():
        raise ValueError(f"No se encontró la plantilla TDR: {ruta_plantilla}")

    valores_texto = {campo.name: _texto_valor(campo) for campo in campos if campo.name}
    valores_html = {
        campo.name: campo.value
        for campo in campos
        if campo.name
        and not campo.value_text
        and isinstance(campo.value, str)
        and _PATRON_HTML.search(campo.value)
    }

    documento = Document(str(ruta_plantilla))

    for parrafo in documento.paragraphs:
        _reemplazar_en_parrafo(parrafo, valores_texto, valores_html)

    for tabla in _iterar_tablas(documento.tables):
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    _reemplazar_en_parrafo(parrafo, valores_texto, valores_html)

    directorio_salida = Path(STORAGE_PATH, "soportes", "tdr", "priview_tdr")
    directorio_salida.mkdir(parents=True, exist_ok=True)

    docx_temporal = directorio_salida / f"{identificador}.docx"
    documento.save(str(docx_temporal))
    try:
        ruta_pdf_final = _convertir_a_pdf(docx_temporal)
    finally:
        docx_temporal.unlink(missing_ok=True)

    return ruta_pdf_final


def _texto_valor(campo: TdrFormField) -> str:
    if campo.value_text:
        return _html_a_texto(str(campo.value_text))
    if campo.value is None:
        return ""
    if isinstance(campo.value, bool):
        return "Sí" if campo.value else "No"
    return _html_a_texto(str(campo.value))


def _html_a_texto(valor: str) -> str:
    texto = re.sub(r"(?i)<li[^>]*>", "\n- ", valor)
    texto = re.sub(r"(?i)</li>", "", texto)
    texto = re.sub(r"(?i)<br\s*/?>", "\n", texto)
    texto = re.sub(r"(?i)</(p|div|ul|ol)>", "\n", texto)
    texto = re.sub(r"(?i)<[^>]+>", "", texto)
    texto = unescape(texto)

    lineas = [linea.strip() for linea in texto.splitlines()]
    return "\n".join(linea for linea in lineas if linea)


def _reemplazar_marcadores(texto: str, valores: dict[str, str]) -> str:
    return _PATRON_MARCADOR.sub(lambda m: valores.get(m.group(1), m.group(0)), texto)


def _reemplazar_en_parrafo(paragraph, valores_texto: dict[str, str], valores_html: dict[str, str]) -> None:
    texto_completo = "".join(run.text for run in paragraph.runs)
    if "<<" not in texto_completo:
        return

    marcadores = list(_PATRON_MARCADOR.finditer(texto_completo))
    if len(marcadores) == 1 and texto_completo.strip() == marcadores[0].group(0):
        nombre_campo = marcadores[0].group(1)
        if nombre_campo in valores_html:
            _insertar_html_como_parrafos(paragraph, valores_html[nombre_campo], paragraph.runs[0])
            return

    nuevo_texto = _reemplazar_marcadores(texto_completo, valores_texto)
    if nuevo_texto == texto_completo:
        return

    for run in paragraph.runs[1:]:
        run.text = ""

    if not paragraph.runs:
        paragraph.add_run(nuevo_texto)
        return

    lineas = nuevo_texto.split("\n")
    primer_run = paragraph.runs[0]
    primer_run.text = lineas[0]
    # Los saltos de línea dentro de un mismo run requieren <w:br/> explícito.
    for linea in lineas[1:]:
        primer_run._r.add_br()
        primer_run._r.add_t(linea)


class _RunHtml:
    __slots__ = ("texto", "negrita", "cursiva", "subrayado")

    def __init__(self, texto: str, negrita: bool = False, cursiva: bool = False, subrayado: bool = False):
        self.texto = texto
        self.negrita = negrita
        self.cursiva = cursiva
        self.subrayado = subrayado


class _BloqueHtml:
    def __init__(self, tipo_lista: str | None = None, numero: int | None = None):
        self.runs: list[_RunHtml] = []
        self.tipo_lista = tipo_lista  # None, "ul" o "ol"
        self.numero = numero


class _ParserHtmlWysiwyg(HTMLParser):
    _ETIQUETAS_NEGRITA = {"b", "strong"}
    _ETIQUETAS_CURSIVA = {"i", "em"}
    _ETIQUETAS_SUBRAYADO = {"u"}
    _ETIQUETAS_BLOQUE = {"p", "div", "li"}
    _ETIQUETAS_LISTA = {"ul", "ol"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.bloques: list[_BloqueHtml] = [_BloqueHtml()]
        self._negrita = 0
        self._cursiva = 0
        self._subrayado = 0
        self._pila_listas: list[str] = []
        self._contadores: list[int] = []

    def _bloque_actual(self) -> _BloqueHtml:
        return self.bloques[-1]

    def _nuevo_bloque(self, tipo_lista: str | None = None, numero: int | None = None) -> None:
        if self._bloque_actual().runs:
            self.bloques.append(_BloqueHtml(tipo_lista, numero))
        else:
            self._bloque_actual().tipo_lista = tipo_lista
            self._bloque_actual().numero = numero

    def handle_starttag(self, tag, attrs):
        if tag in self._ETIQUETAS_NEGRITA:
            self._negrita += 1
        elif tag in self._ETIQUETAS_CURSIVA:
            self._cursiva += 1
        elif tag in self._ETIQUETAS_SUBRAYADO:
            self._subrayado += 1
        elif tag == "br":
            self._bloque_actual().runs.append(_RunHtml("\n"))
        elif tag in self._ETIQUETAS_LISTA:
            self._pila_listas.append(tag)
            self._contadores.append(0)
        elif tag == "li":
            tipo_lista = self._pila_listas[-1] if self._pila_listas else None
            numero = None
            if tipo_lista == "ol":
                self._contadores[-1] += 1
                numero = self._contadores[-1]
            self._nuevo_bloque(tipo_lista, numero)
        elif tag in ("p", "div"):
            self._nuevo_bloque()

    def handle_startendtag(self, tag, attrs):
        if tag == "br":
            self._bloque_actual().runs.append(_RunHtml("\n"))

    def handle_endtag(self, tag):
        if tag in self._ETIQUETAS_NEGRITA:
            self._negrita = max(0, self._negrita - 1)
        elif tag in self._ETIQUETAS_CURSIVA:
            self._cursiva = max(0, self._cursiva - 1)
        elif tag in self._ETIQUETAS_SUBRAYADO:
            self._subrayado = max(0, self._subrayado - 1)
        elif tag in self._ETIQUETAS_LISTA:
            if self._pila_listas:
                self._pila_listas.pop()
                self._contadores.pop()
        elif tag in self._ETIQUETAS_BLOQUE:
            self._nuevo_bloque()

    def handle_data(self, data):
        if not data:
            return
        self._bloque_actual().runs.append(
            _RunHtml(data, self._negrita > 0, self._cursiva > 0, self._subrayado > 0)
        )


def _parsear_html_a_bloques(html: str) -> list[_BloqueHtml]:
    parser = _ParserHtmlWysiwyg()
    parser.feed(html)
    return [bloque for bloque in parser.bloques if "".join(r.texto for r in bloque.runs).strip()]


def _insertar_html_como_parrafos(paragraph, html: str, run_base) -> None:
    bloques = _parsear_html_a_bloques(html)

    for bloque in bloques:
        nuevo_parrafo = paragraph.insert_paragraph_before()
        if bloque.tipo_lista == "ol":
            if not _asignar_estilo_lista(nuevo_parrafo, "List Number"):
                _aplicar_fuente(nuevo_parrafo.add_run(f"{bloque.numero}. "), run_base)
        elif bloque.tipo_lista == "ul":
            if not _asignar_estilo_lista(nuevo_parrafo, "List Bullet"):
                _aplicar_fuente(nuevo_parrafo.add_run("• "), run_base)

        for run_html in bloque.runs:
            partes = run_html.texto.split("\n")
            for indice, parte in enumerate(partes):
                if indice > 0:
                    nuevo_parrafo.add_run().add_break()
                if not parte:
                    continue
                run = nuevo_parrafo.add_run(parte)
                _aplicar_fuente(run, run_base)
                run.bold = run_html.negrita or None
                run.italic = run_html.cursiva or None
                run.underline = run_html.subrayado or None

    paragraph._element.getparent().remove(paragraph._element)


def _aplicar_fuente(run, run_base) -> None:
    # Los runs nuevos no heredan la fuente de la plantilla; se copia explícitamente.
    run.font.name = run_base.font.name
    run.font.size = run_base.font.size
    run.font.color.rgb = run_base.font.color.rgb


def _asignar_estilo_lista(paragraph, nombre_estilo: str) -> bool:
    try:
        paragraph.style = nombre_estilo
        return True
    except KeyError:
        return False




def _iterar_tablas(tablas):
    for tabla in tablas:
        yield tabla
        for fila in tabla.rows:
            for celda in fila.cells:
                yield from _iterar_tablas(celda.tables)


def _convertir_a_pdf(archivo: Path) -> Path:
    libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
    if libreoffice is None:
        raise ValueError(
            "La generación del PDF del TDR requiere LibreOffice. "
            "Instale libreoffice-writer en el servidor."
        )

    try:
        subprocess.run(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation={archivo.parent.joinpath('perfil-libreoffice').as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(archivo.parent),
                str(archivo),
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=60,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as error:
        raise ValueError(f"No fue posible generar el PDF del TDR: {error}") from error

    archivo_pdf = archivo.with_suffix(".pdf")
    if not archivo_pdf.exists():
        raise ValueError("LibreOffice no generó el archivo PDF del TDR")

    return archivo_pdf
